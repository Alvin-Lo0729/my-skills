#!/usr/bin/env python3
"""
建立格式化的繁體中文 Word 重點摘要文件。

使用方式：
  python create_docx.py --title "標題" --summary "摘要文字" \
    --points '[{"title": "重點1", "details": ["說明A", "說明B"]}, ...]' \
    --source "來源URL或檔名" --output "output.docx"

或從 Python 呼叫：
  from create_docx import create_summary_docx
  create_summary_docx(title, summary, points, source, output_path)
"""

import argparse
import json
import sys
from datetime import date
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("錯誤：需要安裝 python-docx")
    print("請執行：pip install python-docx")
    sys.exit(1)


def set_font(run, size=12, bold=False, color=None, font_name="Microsoft JhengHei"):
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    run.font.name = font_name
    # 設定東亞字型
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.insert(0, rFonts)


def add_horizontal_line(doc):
    """在段落後加入水平分隔線"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def create_summary_docx(
    title: str,
    summary: str,
    points: list[dict],
    source: str = "",
    output_path: str = None,
) -> str:
    """
    建立 Word 摘要文件。

    Args:
        title: 文件標題
        summary: 整體摘要（2-4 句）
        points: 重點列表，每項格式：
                {"title": "重點標題", "details": ["說明1", "說明2"]}
        source: 來源 URL 或檔名
        output_path: 輸出路徑（預設為桌面）

    Returns:
        輸出檔案的絕對路徑
    """
    doc = Document()

    # 設定頁面邊距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    # ── 標題 ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(6)
    title_run = title_para.add_run(title)
    set_font(title_run, size=20, bold=True)

    # ── 來源與日期（灰色小字）──
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_para.paragraph_format.space_after = Pt(4)

    if source:
        src_run = meta_para.add_run(f"來源：{source}")
        set_font(src_run, size=9, color=(128, 128, 128))
        meta_para.add_run("　")

    date_run = meta_para.add_run(f"整理日期：{date.today().strftime('%Y-%m-%d')}")
    set_font(date_run, size=9, color=(128, 128, 128))

    add_horizontal_line(doc)

    # ── 摘要 ──
    section_para = doc.add_paragraph()
    section_para.paragraph_format.space_before = Pt(10)
    section_para.paragraph_format.space_after = Pt(4)
    section_run = section_para.add_run("【摘要】")
    set_font(section_run, size=14, bold=True)

    summary_para = doc.add_paragraph()
    summary_para.paragraph_format.space_after = Pt(12)
    summary_run = summary_para.add_run(summary)
    set_font(summary_run, size=11)

    add_horizontal_line(doc)

    # ── 重點整理 ──
    points_header = doc.add_paragraph()
    points_header.paragraph_format.space_before = Pt(10)
    points_header.paragraph_format.space_after = Pt(8)
    points_run = points_header.add_run("【重點整理】")
    set_font(points_run, size=14, bold=True)

    for i, point in enumerate(points, 1):
        # 重點標題
        point_para = doc.add_paragraph()
        point_para.paragraph_format.space_before = Pt(8)
        point_para.paragraph_format.space_after = Pt(3)
        point_run = point_para.add_run(f"{i}. {point['title']}")
        set_font(point_run, size=12, bold=True)

        # 子項目說明
        for detail in point.get("details", []):
            detail_para = doc.add_paragraph()
            detail_para.paragraph_format.left_indent = Cm(1.2)
            detail_para.paragraph_format.space_before = Pt(2)
            detail_para.paragraph_format.space_after = Pt(2)
            bullet_run = detail_para.add_run("• ")
            set_font(bullet_run, size=11, color=(80, 80, 80))
            detail_run = detail_para.add_run(detail)
            set_font(detail_run, size=11)

    # 決定輸出路徑
    if not output_path:
        desktop = Path.home() / "Desktop"
        safe_title = "".join(c for c in title[:15] if c.isalnum() or c in " _-（）()【】")
        safe_title = safe_title.strip() or "摘要"
        filename = f"{date.today().strftime('%Y-%m-%d')}_{safe_title}.docx"
        output_path = str(desktop / filename)

    doc.save(output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(description="建立繁體中文重點摘要 Word 文件")
    parser.add_argument("--title", required=True, help="文件標題")
    parser.add_argument("--summary", required=True, help="整體摘要文字")
    parser.add_argument("--points", required=True, help='重點 JSON，格式：[{"title":"...","details":["..."]}]')
    parser.add_argument("--source", default="", help="來源 URL 或檔名")
    parser.add_argument("--output", default=None, help="輸出路徑（預設存到桌面）")
    args = parser.parse_args()

    try:
        points = json.loads(args.points)
    except json.JSONDecodeError as e:
        print(f"錯誤：points 格式不正確 - {e}")
        sys.exit(1)

    output = create_summary_docx(
        title=args.title,
        summary=args.summary,
        points=points,
        source=args.source,
        output_path=args.output,
    )
    print(f"✓ 已儲存：{output}")


if __name__ == "__main__":
    main()
