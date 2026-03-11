#!/usr/bin/env python3
"""
建立雙語對照 Word 翻譯文件。

使用方式：
  python create_translation_docx.py \
    --title "文件標題" \
    --segments '[{"original":"原文","translation":"譯文"}]' \
    --source "來源說明" \
    --output "output.docx"
"""

import argparse
import json
import sys
from datetime import date
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("錯誤：需要安裝 python-docx")
    print("請執行：pip install python-docx")
    sys.exit(1)


def set_font(run, size=12, bold=False, color=None, font_name="Microsoft JhengHei"):
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    run.font.name = font_name
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.insert(0, rFonts)


def add_horizontal_line(doc, color="CCCCCC", thickness="6"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), thickness)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def create_translation_docx(
    title: str,
    segments: list[dict],
    source: str = "",
    output_path: str = None,
) -> str:
    """
    建立雙語對照 Word 翻譯文件。

    Args:
        title: 文件標題
        segments: 段落列表，每項格式：
                  {"original": "原文段落", "translation": "譯文段落"}
        source: 來源說明
        output_path: 輸出路徑（預設為桌面）

    Returns:
        輸出檔案的絕對路徑
    """
    doc = Document()

    # 設定頁面邊距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    # ── 標題 ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(4)
    title_run = title_para.add_run(title)
    set_font(title_run, size=20, bold=True)

    # ── 來源與日期 ──
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_para.paragraph_format.space_after = Pt(4)

    if source:
        src_run = meta_para.add_run(f"來源：{source}")
        set_font(src_run, size=9, color=(128, 128, 128))
        meta_para.add_run("　")

    date_run = meta_para.add_run(f"翻譯日期：{date.today().strftime('%Y-%m-%d')}")
    set_font(date_run, size=9, color=(128, 128, 128))

    add_horizontal_line(doc, color="333333", thickness="12")

    # ── 雙語對照段落 ──
    for i, seg in enumerate(segments):
        original = seg.get("original", "").strip()
        translation = seg.get("translation", "").strip()

        if not original and not translation:
            continue

        # 原文標籤
        orig_label = doc.add_paragraph()
        orig_label.paragraph_format.space_before = Pt(10)
        orig_label.paragraph_format.space_after = Pt(2)
        orig_label_run = orig_label.add_run("【原文】")
        set_font(orig_label_run, size=10, bold=True, color=(100, 100, 100))

        # 原文內容
        orig_para = doc.add_paragraph()
        orig_para.paragraph_format.space_after = Pt(6)
        orig_para.paragraph_format.left_indent = Cm(0.5)
        orig_run = orig_para.add_run(original)
        set_font(orig_run, size=11, color=(60, 60, 60))

        # 譯文標籤
        trans_label = doc.add_paragraph()
        trans_label.paragraph_format.space_before = Pt(4)
        trans_label.paragraph_format.space_after = Pt(2)
        trans_label_run = trans_label.add_run("【譯文】")
        set_font(trans_label_run, size=10, bold=True, color=(0, 80, 160))

        # 譯文內容
        trans_para = doc.add_paragraph()
        trans_para.paragraph_format.space_after = Pt(6)
        trans_para.paragraph_format.left_indent = Cm(0.5)
        trans_run = trans_para.add_run(translation)
        set_font(trans_run, size=11)

        # 段落之間的分隔線（最後一段不加）
        if i < len(segments) - 1:
            add_horizontal_line(doc, color="DDDDDD", thickness="4")

    # 決定輸出路徑
    if not output_path:
        desktop = Path.home() / "Desktop"
        safe_title = "".join(c for c in title[:15] if c.isalnum() or c in " _-（）()【】")
        safe_title = safe_title.strip() or "translation"
        filename = f"{date.today().strftime('%Y-%m-%d')}_translation_{safe_title}.docx"
        output_path = str(desktop / filename)

    doc.save(output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(description="建立雙語對照 Word 翻譯文件")
    parser.add_argument("--title", required=True, help="文件標題")
    parser.add_argument("--segments", required=True,
                        help='段落 JSON，格式：[{"original":"...","translation":"..."}]')
    parser.add_argument("--source", default="", help="來源說明")
    parser.add_argument("--output", default=None, help="輸出路徑（預設存到桌面）")
    args = parser.parse_args()

    try:
        segments = json.loads(args.segments)
    except json.JSONDecodeError as e:
        print(f"錯誤：segments 格式不正確 - {e}")
        sys.exit(1)

    output = create_translation_docx(
        title=args.title,
        segments=segments,
        source=args.source,
        output_path=args.output,
    )
    print(f"✓ 已儲存：{output}")


if __name__ == "__main__":
    main()
